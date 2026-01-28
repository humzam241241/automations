"""
Attachment Reader - Extracts text from various file formats.
Supports: PDF, Word (.docx), Text files, and more.
"""
import os
import re
from pathlib import Path
from typing import Optional


def extract_text_from_file(file_path: str) -> str:
    """
    Extract text from a file based on its extension.
    
    Supports:
    - PDF files (.pdf)
    - Word documents (.docx)
    - Text files (.txt)
    - CSV files (.csv)
    - Excel files (.xlsx, .xls)
    
    Args:
        file_path: Path to the file
    
    Returns:
        Extracted text content
    """
    if not file_path or not os.path.exists(file_path):
        return ""
    
    ext = Path(file_path).suffix.lower()
    
    try:
        if ext == '.pdf':
            return extract_pdf_text(file_path)
        elif ext == '.docx':
            return extract_docx_text(file_path)
        elif ext == '.doc':
            return extract_doc_text(file_path)
        elif ext in ['.xlsx', '.xls']:
            return extract_excel_text(file_path)
        elif ext in ['.txt', '.csv', '.log', '.md']:
            return extract_plain_text(file_path)
        else:
            return ""
    except Exception as e:
        print(f"Warning: Could not extract text from {file_path}: {e}")
        return ""


def extract_pdf_text(file_path: str) -> str:
    """Extract text from a PDF file."""
    try:
        from PyPDF2 import PdfReader
        
        reader = PdfReader(file_path)
        text_parts = []
        
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        
        return "\n\n".join(text_parts)
    except ImportError:
        print("PyPDF2 not installed. Install with: pip install PyPDF2")
        return ""
    except Exception as e:
        print(f"PDF extraction error: {e}")
        return ""


def extract_docx_text(file_path: str) -> str:
    """Extract text from a Word document (.docx)."""
    try:
        from docx import Document
        
        doc = Document(file_path)
        text_parts = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)
        
        return "\n".join(text_parts)
    except ImportError:
        print("python-docx not installed. Install with: pip install python-docx")
        return ""
    except Exception as e:
        print(f"DOCX extraction error: {e}")
        return ""


def extract_doc_text(file_path: str) -> str:
    """Extract text from old Word format (.doc) - limited support."""
    # .doc files require additional tools (antiword, etc.)
    # For now, return empty and log warning
    print(f"Warning: .doc format not fully supported: {file_path}")
    return ""


def extract_excel_text(file_path: str) -> str:
    """
    Extract text from Excel files (.xlsx, .xls).
    Converts tabular data to readable text with clear column:value format.
    """
    try:
        from openpyxl import load_workbook
        
        wb = load_workbook(file_path, read_only=True, data_only=True)
        text_parts = []
        
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = list(ws.iter_rows(values_only=True))
            
            if not rows:
                continue
            
            # First row as headers
            headers = [str(h).strip() if h else f"Col{i}" for i, h in enumerate(rows[0])]
            
            text_parts.append(f"=== Sheet: {sheet_name} ===")
            text_parts.append("Headers: " + " | ".join(headers))
            
            # Each data row - format as "Header: Value" pairs for better extraction
            for row_idx, row in enumerate(rows[1:], start=1):
                row_text_parts = []
                for i, val in enumerate(row):
                    if val is not None and str(val).strip():
                        header = headers[i] if i < len(headers) else f"Col{i}"
                        row_text_parts.append(f"{header}: {val}")
                
                if row_text_parts:
                    text_parts.append(f"Row {row_idx}: " + " | ".join(row_text_parts))
        
        wb.close()
        return "\n".join(text_parts)
    
    except ImportError:
        print("openpyxl not installed. Install with: pip install openpyxl")
        return ""
    except Exception as e:
        print(f"Excel extraction error: {e}")
        return ""


def extract_excel_rows(file_path: str) -> list:
    """
    Extract Excel data as list of dicts (one dict per row).
    This is used when we want to EXPAND an email with Excel attachment
    into multiple records.
    
    Returns:
        List of dicts, each dict represents a row with header:value pairs
    """
    try:
        from openpyxl import load_workbook
        
        wb = load_workbook(file_path, read_only=True, data_only=True)
        all_rows = []
        
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = list(ws.iter_rows(values_only=True))
            
            if len(rows) < 2:  # Need headers + at least 1 data row
                continue
            
            # First row as headers
            headers = [str(h).strip() if h else f"Col{i}" for i, h in enumerate(rows[0])]
            
            # Each data row becomes a dict
            for row in rows[1:]:
                row_dict = {"_sheet": sheet_name}
                has_data = False
                
                for i, val in enumerate(row):
                    header = headers[i] if i < len(headers) else f"Col{i}"
                    if val is not None:
                        row_dict[header] = str(val).strip()
                        if str(val).strip():
                            has_data = True
                    else:
                        row_dict[header] = ""
                
                if has_data:  # Only include rows with at least some data
                    all_rows.append(row_dict)
        
        wb.close()
        return all_rows
    
    except ImportError:
        return []
    except Exception as e:
        print(f"Excel row extraction error: {e}")
        return []


def extract_excel_rows_from_bytes(data: bytes) -> list:
    """Extract Excel rows from bytes (for email attachments)."""
    import tempfile
    
    try:
        with tempfile.NamedTemporaryFile(suffix='.xlsx', delete=False) as tmp:
            tmp.write(data)
            tmp_path = tmp.name
        
        rows = extract_excel_rows(tmp_path)
        
        try:
            os.unlink(tmp_path)
        except:
            pass
        
        return rows
    except Exception as e:
        print(f"Excel bytes extraction error: {e}")
        return []


def extract_plain_text(file_path: str) -> str:
    """Extract text from a plain text file."""
    try:
        encodings = ['utf-8', 'latin-1', 'cp1252', 'ascii']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        
        # Fallback: binary read and decode what we can
        with open(file_path, 'rb') as f:
            return f.read().decode('utf-8', errors='ignore')
    except Exception as e:
        print(f"Text file read error: {e}")
        return ""


def extract_attachment_from_bytes(data: bytes, filename: str) -> str:
    """
    Extract text from attachment data in memory.
    
    Args:
        data: Raw bytes of the attachment
        filename: Original filename (for extension detection)
    
    Returns:
        Extracted text
    """
    import tempfile
    
    ext = Path(filename).suffix.lower()
    
    # Supported file types for text extraction
    if ext not in ['.pdf', '.docx', '.txt', '.csv', '.xlsx', '.xls']:
        return ""
    
    try:
        # Write to temp file and extract
        with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp:
            tmp.write(data)
            tmp_path = tmp.name
        
        text = extract_text_from_file(tmp_path)
        
        # Clean up
        try:
            os.unlink(tmp_path)
        except:
            pass
        
        return text
    except Exception as e:
        print(f"Attachment extraction error: {e}")
        return ""


def extract_numbers_from_text(text: str) -> list:
    """
    Extract all reference numbers and IDs from text.
    
    Returns list of potential reference numbers found.
    """
    patterns = [
        r'[A-Z]{2,5}[-#]?\d{3,10}',  # MI-12345, CAPA123
        r'\d{5,10}',  # Long numbers
        r'[A-Z]{2,3}\d{2}[-/]\d{4,6}',  # PO22-123456
    ]
    
    numbers = []
    for pattern in patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        numbers.extend(matches)
    
    return list(set(numbers))
